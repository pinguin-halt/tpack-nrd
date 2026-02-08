#!/usr/bin/env python3
"""General Markdown to DOCX converter for academic manuscripts.

Defaults are configured for this repository so terminal commands stay short.
"""

from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path


REPO_ROOT = Path(__file__).resolve().parents[3]

DEFAULTS = {
    "id": {
        "input": "paper/paper_full_id.md",
        "output": "paper/word_pipeline/docx/paper_full_id_final.docx",
    },
    "en": {
        "input": "paper/paper_full_en.md",
        "output": "paper/word_pipeline/docx/paper_full_en_final.docx",
    },
}

DEFAULT_TEMPLATE = "paper/word_pipeline/template/elsevier_smart_health_template.docx"
DEFAULT_BIBLIOGRAPHY_CANDIDATES = [
    "paper/references.bib",
    "references.bib",
]
DEFAULT_CSL_CANDIDATES = [
    "paper/apa.csl",
    "apa.csl",
]


def to_abs(path_value: str | None) -> Path | None:
    if not path_value:
        return None
    path = Path(path_value)
    if not path.is_absolute():
        path = REPO_ROOT / path
    return path


def first_existing(candidates: list[str]) -> Path | None:
    for value in candidates:
        path = to_abs(value)
        if path and path.exists():
            return path
    return None


def require_pandoc() -> str:
    pandoc_path = shutil.which("pandoc")
    if pandoc_path:
        return pandoc_path

    print(
        "Error: pandoc is not installed.\n"
        "Install one of these:\n"
        "- conda run -n tpack-research conda install -y -n tpack-research -c conda-forge pandoc\n"
        "- brew install pandoc",
        file=sys.stderr,
    )
    sys.exit(1)


def has_manual_heading_numbering(markdown_path: Path) -> bool:
    pattern = r"^\s*#{1,6}\s+\d+(?:\.\d+)*\.?\s+"
    try:
        text = markdown_path.read_text(encoding="utf-8", errors="ignore")
    except OSError:
        return False

    for line in text.splitlines():
        if re.match(pattern, line):
            return True
    return False


def create_heading_normalized_copy(
    markdown_path: Path,
    unnumber_frontmatter: bool,
    shift_heading_levels: bool = True,  # kept for API compat; always True
) -> Path | None:
    """Create a temp markdown copy with heading numbers stripped and all
    headings at level >= 2 shifted up by one so Pandoc ``--number-sections``
    produces correct output.

    Logic:
    - The document title (first ``#`` heading) stays at level 1, tagged
      ``{.unnumbered}``.
    - All other headings at level >= 2 are shifted up by one (``##`` → ``#``,
      ``###`` → ``##``, etc.) so that main body sections become Pandoc level 1.
    - Headings with a manual numeric prefix (e.g. ``## 2. Foo``) have that
      prefix stripped so Pandoc auto-numbers them.
    - Headings **without** a numeric prefix get ``{.unnumbered}`` so Pandoc
      skips numbering them.

    Note: *shift_heading_levels* is accepted for backward compatibility but
    the shift is now always applied (the non-shift approach produced
    inconsistent hierarchies).
    """
    pattern_numbered = re.compile(r"^(\s*#{1,6}\s+)\d+(?:\.\d+)*\.?\s+(.*)$")
    pattern_heading = re.compile(r"^(\s*#{1,6}\s+)(.*)$")
    try:
        original = markdown_path.read_text(encoding="utf-8", errors="ignore")
    except OSError:
        return None

    changed = False
    normalized_lines: list[str] = []
    title_seen = False  # track the very first heading (document title)

    for line in original.splitlines():
        # --- Check if this line is a heading at all ---
        match_heading = pattern_heading.match(line)
        if not match_heading:
            normalized_lines.append(line)
            continue

        raw_hashes = match_heading.group(1)
        level = raw_hashes.count("#")

        # --- Document title: first # heading, never numbered, never shifted ---
        if not title_seen and level == 1:
            title_seen = True
            heading_text = match_heading.group(2).strip()
            if "{.unnumbered}" not in heading_text:
                normalized_lines.append(f"# {heading_text} {{.unnumbered}}")
                changed = True
            else:
                normalized_lines.append(line)
            continue

        # --- All headings at level >= 2: shift up by one ---
        new_level = max(level - 1, 1)

        # --- Check if heading has a manual numeric prefix ---
        match_numbered = pattern_numbered.match(line)

        if match_numbered:
            # Numbered heading: strip prefix, shift level, let Pandoc number it
            heading_text = match_numbered.group(2).strip()
            normalized_lines.append(f"{'#' * new_level} {heading_text}")
            changed = True
        else:
            # Non-numbered heading: shift level AND tag as {.unnumbered}
            heading_text = match_heading.group(2).strip()

            if unnumber_frontmatter and "{.unnumbered}" not in heading_text:
                normalized_lines.append(
                    f"{'#' * new_level} {heading_text} {{.unnumbered}}"
                )
                changed = True
            else:
                # Already has {.unnumbered} but still needs level shift
                normalized_lines.append(f"{'#' * new_level} {heading_text}")
                if new_level != level:
                    changed = True

    if not changed:
        return None

    temp = tempfile.NamedTemporaryFile(
        mode="w",
        encoding="utf-8",
        suffix=".md",
        prefix="_heading_normalized_",
        dir=str(markdown_path.parent),
        delete=False,
    )
    with temp:
        temp.write("\n".join(normalized_lines) + "\n")

    return Path(temp.name)


def _set_cell_borders(cell: object, qn: object, OxmlElement: object) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right"):
        edge_tag = f"w:{edge}"
        edge_el = tc_borders.find(qn(edge_tag))
        if edge_el is None:
            edge_el = OxmlElement(edge_tag)
            tc_borders.append(edge_el)
        edge_el.set(qn("w:val"), "single")
        edge_el.set(qn("w:sz"), "8")
        edge_el.set(qn("w:space"), "0")
        edge_el.set(qn("w:color"), "000000")


def ensure_template_heading_styles(template_path: Path) -> Path:
    """Return a patched copy of *template_path* that has Heading 3 (and 4/5)
    styles so Pandoc doesn't degrade ``###``/``####`` to Normal.

    If the template already contains Heading 3, return the original path
    unchanged (no temp file created).
    """
    try:
        from docx import Document
    except ImportError:
        return template_path

    doc = Document(str(template_path))
    style_names = {s.name for s in doc.styles}

    missing = [
        n for n in ("Heading 3", "Heading 4", "Heading 5") if n not in style_names
    ]
    if not missing:
        return template_path

    from docx.enum.style import WD_STYLE_TYPE

    # Use Heading 2 as the basis for missing heading styles
    base = None
    for candidate in ("Heading 2", "Heading 1"):
        if candidate in style_names:
            base = doc.styles[candidate]
            break

    for name in missing:
        new_style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        new_style.quick_style = True
        if base and base.font:
            new_style.font.name = base.font.name
            new_style.font.bold = True
            # Slightly smaller for deeper levels
            if base.font.size:
                level = int(name.split()[-1])
                from docx.shared import Pt

                base_pt = base.font.size.pt if base.font.size else 12
                new_style.font.size = Pt(max(base_pt - (level - 2) * 1, 9))

    temp = tempfile.NamedTemporaryFile(
        suffix=".docx",
        prefix="_patched_template_",
        dir=str(template_path.parent),
        delete=False,
    )
    temp.close()
    doc.save(temp.name)
    return Path(temp.name)


def postprocess_docx(
    output_path: Path,
    fix_tables: bool,
    fix_paragraphs: bool,
    els_style_map: bool,
    fix_figures: bool,
) -> None:
    try:
        from docx import Document
        from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        print(
            "Warning: python-docx not installed; skipping table fix.\n"
            "Install it with: conda run -n tpack-research python -m pip install python-docx",
            file=sys.stderr,
        )
        return

    doc = Document(str(output_path))
    style_names = {s.name for s in doc.styles}
    table_para_style_name: str | None = None
    if "Els_Paragraph" in style_names:
        table_para_style_name = "Els_Paragraph"
    elif "Body Text" in style_names:
        table_para_style_name = "Body Text"
    elif "Normal" in style_names:
        table_para_style_name = "Normal"

    if fix_tables:
        for table in doc.tables:
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            try:
                table.style = "Table Grid"
            except Exception:
                pass

            for row_idx, row in enumerate(table.rows):
                for cell in row.cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    _set_cell_borders(cell, qn, OxmlElement)

                    for paragraph in cell.paragraphs:
                        if table_para_style_name:
                            paragraph.style = doc.styles[table_para_style_name]
                        paragraph.alignment = (
                            WD_ALIGN_PARAGRAPH.CENTER
                            if row_idx == 0
                            else WD_ALIGN_PARAGRAPH.LEFT
                        )
                        paragraph.paragraph_format.first_line_indent = None

    if fix_figures:
        max_figure_width_emu = 4_800_000

        for shape in doc.inline_shapes:
            if shape.width > max_figure_width_emu:
                ratio = shape.height / shape.width if shape.width else 1
                shape.width = max_figure_width_emu
                shape.height = int(max_figure_width_emu * ratio)

        for paragraph in doc.paragraphs:
            has_drawing = any(run._r.xpath(".//w:drawing") for run in paragraph.runs)
            if has_drawing:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.first_line_indent = None

    if fix_paragraphs:
        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name if paragraph.style else ""
            text = (paragraph.text or "").strip()

            if not text:
                continue

            if style_name.startswith("Heading") or style_name in {
                "Title",
                "Subtitle",
                "Caption",
                "TOC Heading",
            }:
                continue

            if style_name.startswith("TOC"):
                continue

            if style_name in {"Bibliography", "Reference", "References"}:
                continue

            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Inches(0.25)

    if els_style_map:
        has_els = "Els_Paragraph" in style_names

        if has_els:
            in_references = False
            title_set = False

            for paragraph in doc.paragraphs:
                text = (paragraph.text or "").strip()
                if not text:
                    continue

                style_name = paragraph.style.name if paragraph.style else ""

                lower_text = text.lower()
                if "daftar pustaka" in lower_text or lower_text == "references":
                    if "Els_ReferencesHeading" in style_names:
                        paragraph.style = doc.styles["Els_ReferencesHeading"]
                    in_references = True
                    continue

                if in_references:
                    if "Els_References" in style_names:
                        paragraph.style = doc.styles["Els_References"]
                    continue

                if not title_set and style_name.startswith("Heading"):
                    if "Els_ArticleTitle" in style_names:
                        paragraph.style = doc.styles["Els_ArticleTitle"]
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph.paragraph_format.first_line_indent = None
                    title_set = True
                    continue

                # Map Pandoc heading levels to Elsevier styles.
                # After uniform level-shift (all >= 2 shifted up by 1):
                #   Heading 1 = main sections + frontmatter  -> Els_Heading1
                #   Heading 2 = subsections (2.1, 3.4, etc.) -> Els_Heading2
                #   Heading 3 = sub-subsections (3.4.1, etc.)-> Els_Heading3
                #   Heading 4+ = deeper (rare)               -> Els_Heading3
                if style_name in {"Heading 1", "Heading1"}:
                    if "Els_Heading1" in style_names:
                        paragraph.style = doc.styles["Els_Heading1"]
                    continue

                if style_name in {"Heading 2", "Heading2"}:
                    if "Els_Heading2" in style_names:
                        paragraph.style = doc.styles["Els_Heading2"]
                    continue

                if style_name in {"Heading 3", "Heading3"}:
                    if "Els_Heading3" in style_names:
                        paragraph.style = doc.styles["Els_Heading3"]
                    continue

                if style_name in {
                    "Heading 4",
                    "Heading4",
                    "Heading 5",
                    "Heading5",
                }:
                    target = (
                        "Els_Heading4"
                        if "Els_Heading4" in style_names
                        else "Els_Heading3"
                    )
                    if target in style_names:
                        paragraph.style = doc.styles[target]
                    continue

                if text.startswith("Gambar ") and "Els_FigureCaption" in style_names:
                    paragraph.style = doc.styles["Els_FigureCaption"]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.first_line_indent = None
                    continue

                if text.startswith("Figure ") and "Els_FigureCaption" in style_names:
                    paragraph.style = doc.styles["Els_FigureCaption"]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.first_line_indent = None
                    continue

                if text.startswith("Tabel ") and "Els_TableCaption" in style_names:
                    paragraph.style = doc.styles["Els_TableCaption"]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.first_line_indent = None
                    continue

                if text.startswith("Table ") and "Els_TableCaption" in style_names:
                    paragraph.style = doc.styles["Els_TableCaption"]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.first_line_indent = None
                    continue

                if (
                    style_name in {"Normal", "Body Text"}
                    and "Els_Paragraph" in style_names
                ):
                    paragraph.style = doc.styles["Els_Paragraph"]

    doc.save(str(output_path))


def build_command(args: argparse.Namespace) -> list[str]:
    pandoc = require_pandoc()

    input_path = to_abs(args.input)
    output_path = to_abs(args.output)
    template_path = to_abs(args.template) if args.template else None

    if not input_path or not input_path.exists():
        print(f"Error: input file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    if not output_path:
        print("Error: output path is empty", file=sys.stderr)
        sys.exit(1)

    if template_path and not template_path.exists():
        print(f"Error: template file not found: {template_path}", file=sys.stderr)
        sys.exit(1)

    output_path.parent.mkdir(parents=True, exist_ok=True)

    bibliography = to_abs(args.bibliography) if args.bibliography else None
    if bibliography and not bibliography.exists():
        print(f"Error: bibliography file not found: {bibliography}", file=sys.stderr)
        sys.exit(1)

    csl = to_abs(args.csl) if args.csl else None
    if csl and not csl.exists():
        print(f"Error: CSL file not found: {csl}", file=sys.stderr)
        sys.exit(1)

    resource_path = f"{input_path.parent}:{REPO_ROOT}"

    if args.number_sections is None:
        enable_number_sections = not has_manual_heading_numbering(input_path)
    else:
        enable_number_sections = args.number_sections

    cmd = [
        pandoc,
        str(input_path),
        "--from",
        "markdown+tex_math_dollars+raw_tex+pipe_tables+grid_tables",
        "--to",
        "docx",
        "--standalone",
        "--output",
        str(output_path),
        "--resource-path",
        resource_path,
    ]

    if enable_number_sections:
        cmd.append("--number-sections")

    if args.toc:
        cmd.extend(["--toc", "--toc-depth", str(args.toc_depth)])

    if template_path:
        cmd.extend(["--reference-doc", str(template_path)])

    if bibliography:
        cmd.extend(["--citeproc", "--bibliography", str(bibliography)])

    if csl:
        cmd.extend(["--csl", str(csl)])

    if args.metadata_file:
        metadata_file = to_abs(args.metadata_file)
        if not metadata_file or not metadata_file.exists():
            print(f"Error: metadata file not found: {metadata_file}", file=sys.stderr)
            sys.exit(1)
        cmd.extend(["--metadata-file", str(metadata_file)])

    return cmd


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Convert academic markdown manuscript to journal-ready DOCX."
    )
    parser.add_argument(
        "--lang",
        choices=["id", "en"],
        default="id",
        help="Choose default input/output language profile.",
    )
    parser.add_argument(
        "--input",
        help="Input markdown path (relative to repo root by default).",
    )
    parser.add_argument(
        "--output",
        help="Output docx path (relative to repo root by default).",
    )
    parser.add_argument(
        "--template",
        default=DEFAULT_TEMPLATE,
        help="DOCX template path. Set empty string to skip template.",
    )
    parser.add_argument(
        "--toc",
        action=argparse.BooleanOptionalAction,
        default=False,
        help="Enable/disable table of contents.",
    )
    parser.add_argument(
        "--toc-depth",
        type=int,
        default=3,
        help="TOC depth level.",
    )
    parser.add_argument(
        "--bibliography",
        help="Optional .bib file path. If omitted, auto-detect common files.",
    )
    parser.add_argument(
        "--csl",
        help="Optional CSL style path. If omitted, auto-detect common files.",
    )
    parser.add_argument(
        "--metadata-file",
        help="Optional Pandoc metadata YAML file.",
    )
    parser.add_argument(
        "--number-sections",
        action=argparse.BooleanOptionalAction,
        default=None,
        help=(
            "Enable/disable section numbering. Default auto: disabled if headings "
            "already contain manual numbers."
        ),
    )
    parser.add_argument(
        "--unnumber-frontmatter",
        action=argparse.BooleanOptionalAction,
        default=True,
        help=(
            "Add '{.unnumbered}' to headings without explicit numeric prefixes in "
            "the temporary markdown copy before conversion."
        ),
    )
    parser.add_argument(
        "--shift-heading-levels",
        action=argparse.BooleanOptionalAction,
        default=False,
        help=(
            "Shift markdown headings up by one level in temporary copy "
            "(##->#, ###->##). Usually NOT needed — keep False."
        ),
    )
    parser.add_argument(
        "--normalize-headings",
        action=argparse.BooleanOptionalAction,
        default=True,
        help=(
            "Strip manual numeric prefixes in markdown headings for DOCX conversion "
            "(e.g., '## 2. Intro' -> '## Intro')."
        ),
    )
    parser.add_argument(
        "--fix-tables",
        action=argparse.BooleanOptionalAction,
        default=True,
        help="Post-process DOCX to enforce table borders and alignment.",
    )
    parser.add_argument(
        "--fix-paragraphs",
        action=argparse.BooleanOptionalAction,
        default=True,
        help="Post-process DOCX to enforce justified body paragraphs.",
    )
    parser.add_argument(
        "--els-style-map",
        action=argparse.BooleanOptionalAction,
        default=True,
        help="Map generic styles to Elsevier (Els_*) styles when available.",
    )
    parser.add_argument(
        "--fix-figures",
        action=argparse.BooleanOptionalAction,
        default=True,
        help="Center figures and shrink oversized images in DOCX output.",
    )

    args = parser.parse_args()

    profile = DEFAULTS[args.lang]
    args.input = args.input or profile["input"]
    args.output = args.output or profile["output"]

    if args.template == "":
        args.template = None

    if not args.bibliography:
        autodetected_bib = first_existing(DEFAULT_BIBLIOGRAPHY_CANDIDATES)
        args.bibliography = str(autodetected_bib) if autodetected_bib else None

    if not args.csl:
        autodetected_csl = first_existing(DEFAULT_CSL_CANDIDATES)
        args.csl = str(autodetected_csl) if autodetected_csl else None

    return args


def main() -> None:
    args = parse_args()
    temp_input: Path | None = None
    temp_template: Path | None = None

    try:
        if args.normalize_headings:
            input_abs = to_abs(args.input)
            if input_abs and input_abs.exists():
                temp_input = create_heading_normalized_copy(
                    input_abs,
                    unnumber_frontmatter=args.unnumber_frontmatter,
                    shift_heading_levels=args.shift_heading_levels,
                )
                if temp_input:
                    args.input = str(temp_input)

        # Ensure template has Heading 3/4/5 so Pandoc doesn't degrade them
        if args.template:
            template_abs = to_abs(args.template)
            if template_abs and template_abs.exists():
                patched = ensure_template_heading_styles(template_abs)
                if patched != template_abs:
                    temp_template = patched
                    args.template = str(temp_template)

        cmd = build_command(args)

        print("Running:")
        print(" ".join(cmd))

        result = subprocess.run(cmd, check=False)
        if result.returncode != 0:
            print("Error: pandoc conversion failed.", file=sys.stderr)
            sys.exit(result.returncode)

        if (
            args.fix_tables
            or args.fix_paragraphs
            or args.els_style_map
            or args.fix_figures
        ):
            postprocess_docx(
                to_abs(args.output),
                fix_tables=args.fix_tables,
                fix_paragraphs=args.fix_paragraphs,
                els_style_map=args.els_style_map,
                fix_figures=args.fix_figures,
            )

        print(f"Success: {args.output}")
        if args.toc:
            print(
                "Note: Word may ask to update fields when opening the file because TOC uses field codes."
            )
    finally:
        if temp_input and temp_input.exists():
            temp_input.unlink()
        if temp_template and temp_template.exists():
            temp_template.unlink()


if __name__ == "__main__":
    main()
